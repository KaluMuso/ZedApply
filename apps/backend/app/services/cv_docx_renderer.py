"""Server-side CV DOCX rendering via python-docx."""
from __future__ import annotations

import io
import time
from collections.abc import Callable

from docx import Document
from docx.shared import Pt

from app.schemas.cv_scratch import (
    BuildFromScratchBody,
    ScratchEducation,
    ScratchExperience,
)

_SKILLS_SEPARATOR = " · "


def _contact_line(body: BuildFromScratchBody) -> str:
    bits = [body.basics.phone, body.basics.email, body.basics.location]
    return " · ".join(b.strip() for b in bits if b and b.strip())


def _role_dates(role: ScratchExperience) -> str:
    start = role.start_date.strip()
    end = role.end_date.strip() or "Present"
    parts = [p for p in (start, end) if p]
    return " – ".join(parts)


def _edu_dates(edu: ScratchEducation) -> str:
    parts = [edu.start_date.strip(), edu.end_date.strip()]
    return " – ".join(p for p in parts if p)


def _role_line(role: ScratchExperience) -> str:
    if not role.title.strip() and not role.company.strip():
        return ""
    loc = f" ({role.location.strip()})" if role.location.strip() else ""
    dates = _role_dates(role)
    date_suffix = f" [{dates}]" if dates else ""
    return f"{role.title.strip()}, {role.company.strip()}{loc}{date_suffix}"


def _edu_line(edu: ScratchEducation) -> str:
    if not edu.degree.strip() and not edu.institution.strip():
        return ""
    loc = f" ({edu.location.strip()})" if edu.location.strip() else ""
    dates = _edu_dates(edu)
    date_suffix = f" [{dates}]" if dates else ""
    gpa = f" GPA: {edu.gpa.strip()}" if edu.gpa.strip() else ""
    return f"{edu.degree.strip()}, {edu.institution.strip()}{loc}{date_suffix}{gpa}"


def _add_header(doc: Document, body: BuildFromScratchBody) -> None:
    name = body.basics.full_name.strip() or "CV"
    doc.add_heading(name, level=0)
    if body.basics.headline.strip():
        headline = doc.add_paragraph(body.basics.headline.strip())
        headline.runs[0].bold = True
        headline.runs[0].font.size = Pt(12)
    contact = _contact_line(body)
    if contact:
        contact_para = doc.add_paragraph(contact)
        contact_para.runs[0].font.size = Pt(10)


def _add_summary_section(doc: Document, body: BuildFromScratchBody) -> None:
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(body.summary.strip())


def _add_experience_section(doc: Document, body: BuildFromScratchBody) -> None:
    doc.add_heading("Experience", level=2)
    for role in body.experience:
        line = _role_line(role)
        if not line:
            continue
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(line)
        title_run.bold = True
        bullets = [b.strip() for b in role.achievements if b.strip()]
        if bullets:
            for bullet in bullets:
                doc.add_paragraph(bullet, style="List Bullet")


def _add_education_section(doc: Document, body: BuildFromScratchBody) -> None:
    doc.add_heading("Education", level=2)
    for edu in body.education:
        line = _edu_line(edu)
        if not line:
            continue
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.bold = True


def _add_skills_section(doc: Document, body: BuildFromScratchBody) -> None:
    doc.add_heading("Skills", level=2)
    cleaned = [s.strip() for s in body.skills if s.strip()]
    doc.add_paragraph(_SKILLS_SEPARATOR.join(cleaned))


def _section_builders(body: BuildFromScratchBody) -> list[tuple[str, Callable[[Document], None]]]:
    sections: list[tuple[str, Callable[[Document], None]]] = []
    if body.style.show_summary and body.summary.strip():
        sections.append(("Summary", lambda d: _add_summary_section(d, body)))

    has_experience = any(
        e.title.strip() or e.company.strip() for e in body.experience
    )
    if has_experience:
        sections.append(("Experience", lambda d: _add_experience_section(d, body)))

    has_education = any(
        e.degree.strip() or e.institution.strip() for e in body.education
    )
    if has_education:
        sections.append(("Education", lambda d: _add_education_section(d, body)))

    if any(s.strip() for s in body.skills):
        sections.append(("Skills", lambda d: _add_skills_section(d, body)))

    return sections


def render_cv_docx(body: BuildFromScratchBody) -> tuple[bytes, int]:
    """Render CV to DOCX bytes. Returns (docx_bytes, render_time_ms)."""
    started = time.perf_counter()
    doc = Document()
    _add_header(doc, body)

    sections = _section_builders(body)
    for idx, (_title, add_section) in enumerate(sections):
        if idx > 0:
            doc.add_page_break()
        add_section(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    elapsed_ms = int((time.perf_counter() - started) * 1000)
    return buffer.getvalue(), elapsed_ms


def section_heading_order(body: BuildFromScratchBody) -> list[str]:
    """Ordered section titles included in the export (for tests)."""
    return [title for title, _ in _section_builders(body)]
