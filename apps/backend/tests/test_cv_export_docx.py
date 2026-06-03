"""Tests for POST /cv/export/docx."""
import io
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

FIXTURE_PATH = Path(__file__).resolve().parent / "fixtures" / "cv_export_sample.docx"


@pytest.fixture
def auth_headers():
    from app.core.deps import get_current_user_id
    from main import app

    app.dependency_overrides[get_current_user_id] = lambda: "export-user"
    yield {"Authorization": "Bearer test-token"}
    app.dependency_overrides.pop(get_current_user_id, None)


@pytest.fixture
def sample_export_body():
    return {
        "summary": (
            "ICAZ-qualified accountant with 8+ years in banking and audit. "
            "Experienced in IFRS reporting and stakeholder management."
        ),
        "basics": {
            "full_name": "Chanda Banda",
            "phone": "+260971234567",
            "email": "chanda.banda@email.com",
            "location": "Lusaka, Zambia",
            "headline": "Chartered Accountant · IFRS",
        },
        "experience": [
            {
                "title": "Senior Accountant",
                "company": "ZANACO",
                "location": "Lusaka",
                "start_date": "Jan 2019",
                "end_date": "Present",
                "achievements": [
                    "Led month-end close for 12 branches, reducing cycle time by 18%.",
                ],
            }
        ],
        "education": [
            {
                "degree": "Bachelor of Accountancy",
                "institution": "University of Zambia",
                "location": "Lusaka",
                "start_date": "2011",
                "end_date": "2014",
                "gpa": "Distinction",
            }
        ],
        "skills": ["IFRS", "Excel", "SAP"],
        "style": {"template": "modern", "accent_color": "#0E5C3A", "show_summary": True},
    }


@pytest.fixture
def fixture_docx_bytes() -> bytes:
    """Golden DOCX bytes checked into tests/fixtures (regenerate via test if missing)."""
    if FIXTURE_PATH.is_file():
        return FIXTURE_PATH.read_bytes()
    from app.schemas.cv_scratch import (
        BuildFromScratchBody,
        ScratchBasics,
        ScratchEducation,
        ScratchExperience,
        ScratchStyle,
    )
    from app.services.cv_docx_renderer import render_cv_docx

    body = BuildFromScratchBody(
        summary="Fixture summary for export tests.",
        basics=ScratchBasics(
            full_name="Fixture User",
            phone="+260971234567",
            email="fixture@example.com",
            location="Lusaka",
            headline="Accountant",
        ),
        experience=[
            ScratchExperience(
                title="Analyst",
                company="ACME",
                location="Lusaka",
                start_date="2020",
                end_date="2024",
                achievements=["Delivered reporting."],
            )
        ],
        education=[
            ScratchEducation(
                degree="BSc Finance",
                institution="UNZA",
                location="Lusaka",
                start_date="2016",
                end_date="2019",
                gpa="",
            )
        ],
        skills=["Excel"],
        style=ScratchStyle(),
    )
    docx_bytes, _ = render_cv_docx(body)
    FIXTURE_PATH.parent.mkdir(parents=True, exist_ok=True)
    FIXTURE_PATH.write_bytes(docx_bytes)
    return docx_bytes


def test_export_docx_requires_auth(client: TestClient, sample_export_body):
    res = client.post("/api/v1/cv/export/docx", json=sample_export_body)
    assert res.status_code in (401, 403)


def test_export_docx_returns_docx_bytes(
    client: TestClient, auth_headers, sample_export_body
):
    res = client.post(
        "/api/v1/cv/export/docx",
        json=sample_export_body,
        headers=auth_headers,
    )
    assert res.status_code == 200
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in res.headers.get("content-disposition", "")
    assert "Chanda-Banda.docx" in res.headers["content-disposition"]
    assert res.content[:2] == b"PK"

    doc = Document(io.BytesIO(res.content))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Chanda Banda" in joined
    assert "Summary" in joined
    assert "Senior Accountant" in joined
    assert "University of Zambia" in joined
    assert "IFRS · Excel · SAP" in joined


def test_export_docx_fixture_bytes_are_valid_docx(fixture_docx_bytes: bytes):
    assert fixture_docx_bytes[:2] == b"PK"
    doc = Document(io.BytesIO(fixture_docx_bytes))
    assert any(p.text.strip() == "Fixture User" for p in doc.paragraphs)


def test_export_docx_422_without_full_name(client: TestClient, auth_headers):
    res = client.post(
        "/api/v1/cv/export/docx",
        json={
            "summary": "",
            "basics": {
                "full_name": "   ",
                "phone": "",
                "email": "",
                "location": "",
                "headline": "",
            },
            "experience": [],
            "education": [],
            "skills": [],
            "style": {"show_summary": True},
        },
        headers=auth_headers,
    )
    assert res.status_code == 422
