"""Tests for server-side CV DOCX rendering."""
import io
import zipfile

from docx import Document

from app.schemas.cv_scratch import (
    BuildFromScratchBody,
    ScratchBasics,
    ScratchEducation,
    ScratchExperience,
    ScratchStyle,
)
from app.services.cv_docx_renderer import render_cv_docx, section_heading_order
from tests.test_cv_pdf_renderer import _sample_body


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _count_page_breaks(doc: Document) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'w:type="page"' in run._element.xml:
                count += 1
    return count


def test_render_cv_docx_produces_valid_zip_container():
    docx_bytes, render_ms = render_cv_docx(_sample_body())
    assert docx_bytes[:2] == b"PK"
    assert len(docx_bytes) > 1000
    assert render_ms < 5000
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        assert "[Content_Types].xml" in zf.namelist()
        assert "word/document.xml" in zf.namelist()


def test_render_cv_docx_includes_sections_in_pdf_order():
    body = _sample_body()
    docx_bytes, _ = render_cv_docx(body)
    doc = Document(io.BytesIO(docx_bytes))
    texts = _paragraph_texts(doc)
    assert "Chanda Banda" in texts[0]
    assert section_heading_order(body) == [
        "Summary",
        "Experience",
        "Education",
        "Skills",
    ]
    assert "Summary" in texts
    assert "Senior Accountant" in " ".join(texts)
    assert "University of Zambia" in " ".join(texts)
    assert "IFRS" in " ".join(texts)


def test_render_cv_docx_page_breaks_between_sections():
    body = _sample_body()
    docx_bytes, _ = render_cv_docx(body)
    doc = Document(io.BytesIO(docx_bytes))
    # Four sections → three breaks between them (not before Summary).
    assert _count_page_breaks(doc) == len(section_heading_order(body)) - 1


def test_render_cv_docx_skips_empty_sections():
    body = BuildFromScratchBody(
        summary="",
        basics=ScratchBasics(full_name="Minimal User"),
        experience=[],
        education=[],
        skills=["Excel"],
        style=ScratchStyle(show_summary=False),
    )
    assert section_heading_order(body) == ["Skills"]
